import io
import json

import docx
import pdfplumber

from llm_client import call_llm
from prompts import ADAM_PARSE_SYSTEM


def _strip_fences(raw: str) -> str:
    raw = (raw or "").strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json\n"):
            raw = raw[5:]
    return raw.strip()


def _extract_excel_text(file_bytes: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    priority_keywords = {"variable", "metadata", "codelist", "analysis", "derivation", "dataset", "where"}
    names = wb.sheetnames
    priority = [s for s in names if any(k in s.lower() for k in priority_keywords)]
    rest = [s for s in names if s not in priority]

    lines: list[str] = []
    for s in priority + rest:
        ws = wb[s]
        lines.append(f"\n=== SHEET: {s} ===")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line = " | ".join(cells).strip(" |")
            if line:
                lines.append(line)
    return "\n".join(lines)


def _extract_pdf_text(file_bytes: bytes) -> str:
    lines: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text()
            if txt:
                lines.append(txt)
            for tbl in page.extract_tables() or []:
                for row in tbl:
                    lines.append(" | ".join(str(c) if c is not None else "" for c in (row or [])))
    return "\n".join(lines)


def _extract_docx_text(file_bytes: bytes) -> str:
    d = docx.Document(io.BytesIO(file_bytes))
    lines: list[str] = []
    for p in d.paragraphs:
        if p.text.strip():
            lines.append(p.text)
    for table in d.tables:
        for row in table.rows:
            lines.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(lines)


def parse_adam_specs(
    file_bytes: bytes,
    extension: str,
    provider: str,
    model: str,
    api_key: str,
    parse_temperature: float = 0.4,
) -> dict:
    ext = extension.lower().strip(".")
    if ext in {"xlsx", "xlsm", "xls"}:
        text = _extract_excel_text(file_bytes)
    elif ext == "pdf":
        text = _extract_pdf_text(file_bytes)
    elif ext == "docx":
        text = _extract_docx_text(file_bytes)
    else:
        raise ValueError(f"Unsupported AdaM spec extension: {extension}")

    # Larger cap than v1 to reduce truncation failures on long specs.
    text = text[:120000]
    raw = call_llm(
        system=ADAM_PARSE_SYSTEM,
        user="Extract AdaM specs from this text:\n\n" + text,
        provider=provider,
        model=model,
        api_key=api_key,
        max_tokens=4000,
        temperature=parse_temperature,
        json_mode=True,
    )
    return json.loads(_strip_fences(raw))

