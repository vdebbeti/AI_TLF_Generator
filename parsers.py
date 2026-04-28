import io
import json

import docx
import pdfplumber

from llm_client import call_llm
from prompts import SHELL_PARSE_SYSTEM


def _strip_fences(raw: str) -> str:
    raw = (raw or "").strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json\n"):
            raw = raw[5:]
    return raw.strip()


def _extract_pdf_text(file_bytes: bytes) -> str:
    lines: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for pno, page in enumerate(pdf.pages, start=1):
            lines.append(f"\n--- PAGE {pno} ---")
            tables = page.extract_tables() or []
            for ti, tbl in enumerate(tables, start=1):
                lines.append(f"-- TABLE {ti} --")
                for row in tbl:
                    cooked: list[str] = []
                    for idx, c in enumerate(row or []):
                        txt = str(c) if c is not None else ""
                        txt = txt.rstrip() if idx == 0 else txt.strip()
                        cooked.append(txt)
                    lines.append(" | ".join(cooked))
            text = page.extract_text()
            if text:
                lines.append("-- RAW TEXT --")
                lines.append(text)
    return "\n".join(lines)


def _extract_docx_text(file_bytes: bytes) -> str:
    d = docx.Document(io.BytesIO(file_bytes))
    lines: list[str] = []
    for para in d.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    for ti, table in enumerate(d.tables, start=1):
        lines.append(f"\n--- TABLE {ti} ---")
        for row in table.rows:
            first = row.cells[0]
            indent = ""
            if first.paragraphs:
                pf = first.paragraphs[0].paragraph_format
                if pf and pf.left_indent:
                    indent = "  " * max(1, int(pf.left_indent.pt // 12))
            cells = [c.text.strip() for c in row.cells]
            if cells:
                cells[0] = indent + cells[0]
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def parse_shell(
    file_bytes: bytes,
    extension: str,
    provider: str,
    model: str,
    api_key: str,
    parse_temperature: float = 0.7,
) -> dict:
    ext = extension.lower().strip(".")
    image_exts = {"png", "jpg", "jpeg"}

    if ext in image_exts:
        mime = "image/png" if ext == "png" else "image/jpeg"
        user = (
            "Parse this clinical mock shell image into the required JSON. "
            "Preserve row hierarchy and indentation. "
            "For AE tables, rows are representative examples and must map to SOC/PT dynamic structure."
        )
        raw = call_llm(
            system=SHELL_PARSE_SYSTEM,
            user=user,
            provider=provider,
            model=model,
            api_key=api_key,
            image_bytes=file_bytes,
            image_mime=mime,
            max_tokens=4000,
            temperature=parse_temperature,
        )
    elif ext == "pdf":
        text = _extract_pdf_text(file_bytes)
        raw = call_llm(
            system=SHELL_PARSE_SYSTEM,
            user="Parse this PDF-extracted table shell text:\n\n" + text,
            provider=provider,
            model=model,
            api_key=api_key,
            max_tokens=4000,
            temperature=parse_temperature,
            json_mode=True,
        )
    elif ext == "docx":
        text = _extract_docx_text(file_bytes)
        raw = call_llm(
            system=SHELL_PARSE_SYSTEM,
            user="Parse this DOCX-extracted table shell text:\n\n" + text,
            provider=provider,
            model=model,
            api_key=api_key,
            max_tokens=4000,
            temperature=parse_temperature,
            json_mode=True,
        )
    else:
        raise ValueError(f"Unsupported shell extension: {extension}")

    return json.loads(_strip_fences(raw))

